import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def export_wrongdoing_report(df, output_path="enron_wrongdoing_report.docx"):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    h = doc.add_heading("Top emails — subject and body", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    work = df.reset_index(drop=True)
    sub_col = "original_subject" if "original_subject" in work.columns else "subject"

    for i in range(len(work)):
        row = work.iloc[i]
        doc.add_heading("Email " + str(i + 1), level=1)

        if sub_col in work.columns:
            sj = row[sub_col]
            if pd.isna(sj):
                sj = ""
            subj_line = str(sj)
        else:
            subj_line = ""

        ps = doc.add_paragraph()
        ps.add_run("Subject: ").bold = True
        ps.add_run(subj_line)

        body = ""
        if "text" in work.columns:
            v = row["text"]
            if pd.notna(v):
                body = str(v)
        if not body.strip() and "clean_text" in work.columns:
            v = row["clean_text"]
            if pd.notna(v):
                body = str(v)

        pb = doc.add_paragraph()
        pb.add_run("Body: ").bold = True
        doc.add_paragraph(body if body.strip() else "(no body text)")

    doc.save(output_path)
    return output_path
