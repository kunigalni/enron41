import requests
import pandas as pd
from dateutil import parser
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from preprocessing import preprocess_emails
from keywordFrequency import KeywordFrequency
from sentiment import add_vader_scores
from topic_modeling import run_lda, get_top_words_per_topic, assign_topics
from scoring import add_wrongdoing_score
from time_analysis import compute_time_series, detect_spikes


def fetch_all_emails(host, index, batch_size=5000):
    url = f"{host}{index}/_search?scroll=10m"

    query = {
        "size": batch_size,
        "query": {"match_all": {}}
    }

    r = requests.get(url, json=query, headers={"Content-Type": "application/json"})
    data = r.json()

    scroll_id = data["_scroll_id"]
    hits = data["hits"]["hits"]
    all_hits = hits.copy()

    print(f"[fetch_all_emails] Initial batch: {len(hits)} hits")

    while True:
        scroll_url = f"{host}_search/scroll"
        payload = {"scroll": "10m", "scroll_id": scroll_id}

        r = requests.get(scroll_url, json=payload, headers={"Content-Type": "application/json"})
        data = r.json()

        hits = data["hits"]["hits"]
        if not hits:
            print("[fetch_all_emails] Scroll returned 0 hits — stopping.")
            break

        all_hits.extend(hits)
        scroll_id = data["_scroll_id"]

        if len(all_hits) % 10000 == 0:
            print(f"[fetch_all_emails] Raw hits so far: {len(all_hits)}")

    print(f"[fetch_all_emails] Total raw hits before dedupe: {len(all_hits)}")

    unique = {h["_id"]: h for h in all_hits}
    print(f"[fetch_all_emails] Unique documents after dedupe: {len(unique)}")

    return list(unique.values())


def build_email_dataframe(hits):
    return pd.DataFrame(
        [
            {
                **h["_source"],
                "email_id": h["_id"],
            }
            for h in hits
        ]
    )


def export_top5_word_doc(df, output_path="enron_wrongdoing_report.docx"):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    h = doc.add_heading("Top 5 emails — email_id and full body", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    work = df.reset_index(drop=True)
    sub_col = "original_subject" if "original_subject" in work.columns else "subject"

    for i in range(len(work)):
        row = work.iloc[i]
        doc.add_heading("Email " + str(i + 1), level=1)

        pe = doc.add_paragraph()
        pe.add_run("email_id: ").bold = True
        eid = row["email_id"] if "email_id" in work.columns else ""
        if pd.isna(eid):
            eid = ""
        pe.add_run(str(eid))

        if sub_col in work.columns:
            sj = row[sub_col]
            if pd.isna(sj):
                sj = ""
            subj_line = str(sj)
        else:
            subj_line = ""

        ps = doc.add_paragraph()
        ps.add_run("Subject: ").bold = True
        ps.add_run(subj_line)

        body = ""
        if "full_email_body" in work.columns:
            v = row["full_email_body"]
            if pd.notna(v):
                body = str(v)
        if not body.strip() and "text" in work.columns:
            v = row["text"]
            if pd.notna(v):
                body = str(v)
        if not body.strip() and "clean_text" in work.columns:
            v = row["clean_text"]
            if pd.notna(v):
                body = str(v)

        pb = doc.add_paragraph()
        pb.add_run("Body: ").bold = True
        doc.add_paragraph(body if body.strip() else "(no body text)")

    doc.save(output_path)
    return output_path


def run():
    host = "http://18.188.56.207:9200/"
    index = "enron"
    print(requests.get(host + f"_cat/indices/{index}").content)
    cnt = requests.get(host + f"{index}/_count").json()
    print("[main] count", cnt["count"])
    hits = fetch_all_emails(host, index)
    df = build_email_dataframe(hits)
    df["date"] = df["date"].apply(parser.parse)
    df["date"] = df["date"].dt.tz_localize(None)
    t0 = parser.parse("2001-08-01")
    t1 = parser.parse("2001-12-31")
    df = df[(df["date"] >= t0) & (df["date"] <= t1)].reset_index(drop=True)
    print("[main] rows in window", len(df))

    if "text" in df.columns:
        text_by_email_id = df.set_index("email_id")["text"].to_dict()
    else:
        text_by_email_id = {}

    sub_src = None
    for name in ("subject", "Subject", "SUBJECT"):
        if name in df.columns:
            sub_src = name
            break

    def _orig_subj(val):
        if val is None:
            return ""
        if isinstance(val, float) and pd.isna(val):
            return ""
        s = str(val)
        if s.lower() == "nan":
            return ""
        return s

    if sub_src is not None:
        df["original_subject"] = df[sub_src].map(_orig_subj)
    else:
        df["original_subject"] = ""

    cleaned = preprocess_emails(df.copy(deep=True))
    kf = KeywordFrequency(cleaned)
    kf.compute_tfidf()
    kf.compute_category_counts()
    kf.compute_candidate_scores()

    ts = compute_time_series(kf.df, freq="W")
    ts = detect_spikes(ts, "spe_count")

    candidates = kf.get_candidate_emails(1500)
    candidates = add_vader_scores(candidates)
    lda, topic_matrix, feature_names = run_lda(candidates, text_col="clean_text", n_topics=6)
    topics = get_top_words_per_topic(lda, feature_names)
    candidates = assign_topics(candidates, topic_matrix)
    for tid, words in topics.items():
        print("Topic", tid, ":", ", ".join(words))

    candidates = add_wrongdoing_score(candidates)
    top5 = candidates.sort_values(by="wrongdoing_score", ascending=False).head(5).reset_index(drop=True)

    top5 = top5.copy()
    if "email_id" in top5.columns:
        top5["full_email_body"] = top5["email_id"].map(lambda e: text_by_email_id.get(e, ""))
        top5["full_email_body"] = top5["full_email_body"].apply(
            lambda x: "" if x is None or (isinstance(x, float) and pd.isna(x)) else str(x)
        )
    else:
        top5["full_email_body"] = ""

    sign_cols = [
        "date",
        "email_id",
        "original_subject",
        "candidate_score",
        "wrongdoing_score",
        "spe_score",
        "neg_score",
        "compound_inverse",
        "topic_score",
        "spe_count",
        "risk_term_count",
        "sentiment_neg",
        "sentiment_neu",
        "sentiment_pos",
        "sentiment_compound",
        "dominant_topic",
        "topic_strength",
    ]
    use = [c for c in sign_cols if c in top5.columns]
    print("[main] top 5 — all sign scores")
    with pd.option_context("display.max_columns", None, "display.width", 240, "display.max_colwidth", 50):
        print(top5[use])

    path = export_top5_word_doc(top5, "enron_wrongdoing_report.docx")
    print("[main] wrote", path)


if __name__ == "__main__":
    run()
